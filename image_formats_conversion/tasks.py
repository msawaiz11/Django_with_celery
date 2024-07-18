from celery import shared_task
from django.conf import settings
from PIL import Image
from psd_tools import PSDImage, Group, Layer
from docx import Document
from docx.shared import Inches
import pillow_heif
@shared_task
def jpgtopng(input_image_path, output_image_path):
    img = Image.open(input_image_path)
    img.save(output_image_path, 'PNG')
    return output_image_path

@shared_task
def pngtoword(input_img,output_word):
    doc = Document()

    doc.add_heading('Inserted Image', 0)

    img = Image.open(input_img)
    width_in_inches = 5.0
    height_in_inches = img.height / img.width * width_in_inches

    doc.add_picture(input_img, width=Inches(width_in_inches), height=Inches(height_in_inches))

# Save the document
    doc.save(output_word)

    return output_word


@shared_task 
def imgtopdf(input_path, output_pdf):
    img = Image.open(input_path)
    if img.mode == 'RGBA':
            img = img.convert('RGB')

    img.save(output_pdf, save_all=True)
    
    return output_pdf

@shared_task
def heictoimg(input_img, output_img):
    heif_file = pillow_heif.read_heif(input_img)
    image = Image.frombytes(
        heif_file.mode,
        heif_file.size,
        heif_file.data,
        "raw",
    )
    image.save(output_img, format("png"))

    return output_img